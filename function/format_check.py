import re
import csv
from docx import Document
from docx.shared import Pt 
import xml.etree.ElementTree as ET  
import argparse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# 首页日期格式检测
def check_date_format(doc):
    pattern = r'[〇一二三四五六七八九零]{4}年[〇一二三四五六七八九十]{1,2}月'
    first_page = ''
    for par in doc.paragraphs:
        first_page += par.text + '\n'
        if 'PAGE_BREAK' in par.text:
            break
    match = re.search(pattern, first_page)
    if not match:
        print("首页日期格式有误")

def check_directory_and_font(file_path):
    with open(file_path, 'rb') as file:  
        xml_content = file.read()  
    root = ET.fromstring(xml_content) 
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    for paragraph in root.findall('.//w:p', namespaces=namespaces): 
        texts = [node.text for node in paragraph.findall('.//w:t', namespaces=namespaces) if node.text]
        combined_text = ''.join(texts)
        
        if '目' in combined_text and '录' not in combined_text: 
            print("未找到目录内容") 
        else:  
            rPr = paragraph.find('.//w:rPr', namespaces=namespaces) 
            if rPr is not None:  
                # 检测字体
                rFonts = rPr.find('.//w:rFonts', namespaces=namespaces)
                if rFonts is not None: 
                    font_eastAsia = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')  # 获取东亚字体
                if font_eastAsia != '宋体':  
                    print(f'目录字体为{font_eastAsia},请设置为宋体.') 
            else:
                print('未找到目录字体格式，请检查')
                # 检测字号
                sz = rPr.find('.//w:sz', namespaces=namespaces) 
                szCs = rPr.find('.//w:szCs', namespaces=namespaces) 
                if sz is not None and szCs is not None:  # 如果找到字号元素
                    size = int(sz.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0)) / 2
                    sizeCs = int(szCs.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0)) / 2
    
                    # 检查是否为小四（12磅）
                    if size != 12 or sizeCs != 12:
                        print(f'目录字体大小为 {size} 磅，请设置为小四。')
                else: 
                    print('未设置目录文本字体大小，请检查')

                # 检测行距
                spacing = paragraph.find('.//w:spacing', namespaces=namespaces)
                if spacing is not None: 
                    line_spacing = int(spacing.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line', 0))
                    if line_spacing == 360:
                        print('行距为1.5倍行距。')
                    else:
                        actual_line_spacing = line_spacing / 240.0  # 将行距值转换为倍数
                        print(f'目录行距为{actual_line_spacing:.1f}倍行距，请设置为1.5倍行距。')
                else:  # 如果没有找到行距元素
                    print('No line spacing information found.')  # 输出没有找到行距信息的消息
            break  # 处理完一个段落后跳出循环


# 表名格式检测
def check_table_titles(doc):
    for i, table in enumerate(doc.tables):
        prev_paragraph = None
        for paragraph in doc.paragraphs:
            if table._element == paragraph._element.getnext():
                prev_paragraph = paragraph
                break

        if prev_paragraph:
            prev_paragraph_text = prev_paragraph.text.strip()

            if prev_paragraph_text == "设计文件分发表" or not prev_paragraph_text:
                continue
            prev_paragraph_alignment = prev_paragraph.alignment
            # 检查表名是否左对齐
            is_left_aligned = prev_paragraph_alignment == WD_PARAGRAPH_ALIGNMENT.LEFT
            print(f"表名: '{prev_paragraph_text}'")
            print(f"是否左对齐: {'是' if is_left_aligned else '否'}")

            # 如果左对齐，继续检查表名是否居中
            if is_left_aligned:
                # 假设表名字从第一个空格后开始
                if ' ' in prev_paragraph_text:
                    title_start = prev_paragraph_text.index(' ') + 1
                else:
                    title_start = len(prev_paragraph_text)
                
                # 计算表名前后字符数量
                before_title = title_start
                after_title = len(prev_paragraph_text) - title_start
                
                # 判断是否居中
                char_ratio = before_title / max(after_title, 1)
                if 0.8 <= char_ratio <= 1.2:
                    print("表名字位于居中位置")
                else:
                    print(f"表名字不居中: 前面有 {before_title} 个字符，后面有 {after_title} 个字符")
        else:
            print("未找到与表格相邻的段落")

# 获取对应人员信息
def extract_person_info(text):
    zong_name_match = re.search(r'项目总负责人：([^\n]+)', text)
    dan_name_match = re.search(r'单项设计负责人：([^\n]+)',text)
    jian_name_match = re.search(r'建设单位联系人：([^\n]+)',text)
    tel_match = re.findall(r'电话：([^\n]+)', text)
    mail_match = re.findall(r'电子邮箱：([^\n]+)', text)

    persons = []
    if zong_name_match and len(tel_match)>0 and len(mail_match)>0:
        person1={
            'name':zong_name_match.group(1).strip(),
            'tel':tel_match[0].strip(),
            'mail':mail_match[0].strip()
        }
        persons.append(person1)
    if dan_name_match and len(tel_match) > 1 and len(mail_match) > 1:
        person2={
            'name':dan_name_match.group(1).strip(),
            'tel':tel_match[1].strip(),
            'mail':mail_match[1].strip()
        }   
        persons.append(person2)
    if jian_name_match and len(tel_match) > 2 and len(mail_match) > 2 :
        person3={
            'name':jian_name_match.group(1).strip(),
            'tel':tel_match[2].strip(),
            'mail':mail_match[2].strip()
        }
        persons.append(person3)
    return persons if persons else None


def is_chinese(text):
    for char in text:
        if '\u4e00' <= char <= '\u9fff':
            return True
    return False

def check_paragraph_formatting(doc):
    chinese_font = "宋体"
    english_font = "Times New Roman"
    font_size = Pt(12)
    line_spacing = 1.5
    
    # 确认段落的样式，只针对Normal样式进行检测
    for paragraph in doc.paragraphs:
        # 如果段落内容为空或仅包含空格，则跳过
        if not paragraph.text.strip():
            continue
        
        if paragraph.style.name != 'Normal':
            continue
        
        issues = []
        runs = paragraph.runs
        for run in runs:
            text = run.text.strip()
            if not text:
                continue  # 跳过空文本
            
            if is_chinese(text):
                if run.font.name != chinese_font or run.font.size != font_size:
                    issues.append(f"中文字体应为 {chinese_font}，实际为 {run.font.name if run.font.name else '未设置'}，字体大小应为 {font_size.pt}pt，实际为 {run.font.size.pt if run.font.size else '未设置'}pt")
            else:
                if run.font.name != english_font or run.font.size != font_size:
                    issues.append(f"英文/数字字体应为 {english_font}，实际为 {run.font.name if run.font.name else '未设置'}，字体大小应为 {font_size.pt}pt，实际为 {run.font.size.pt if run.font.size else '未设置'}pt")
        
        if paragraph.paragraph_format.line_spacing != line_spacing:
            issues.append(f"行距应为 {line_spacing} 倍，实际为 {paragraph.paragraph_format.line_spacing if paragraph.paragraph_format.line_spacing else '未设置'}")

        if issues:
            print(f"段落内容: {paragraph.text}\n问题: {', '.join(issues)}\n") 


def find_person_in_csv(csv_path, persons):
    # 读取 CSV 文件
    with open(csv_path, newline='', encoding='GBK') as csvfile:
        reader = csv.DictReader(csvfile)
        
        # 将 CSV 文件中的数据转换为列表
        data = list(reader)
        
    # 查找并输出不完全符合的人员姓名
    unmatched_persons = []
    for person in persons:
        matched = False
        for row in data:
            if (row['姓名'].strip() == person['name'] and
                row['手机号'].strip() == person['tel'] and
                row['邮箱'].strip() == person['mail']):
                matched = True
                break
        if not matched:
            unmatched_persons.append(person['name'])
    
    return unmatched_persons


# 标记找到指定段落后的第一个表格
def find_table_after_paragraph(doc, paragraph_text):
    found_paragraph = False
    
    # 遍历所有段落和表格
    for para in doc.paragraphs:
        if paragraph_text in para.text:
            found_paragraph = True
        
        # 如果已经找到了指定段落，检查下一个表格
        if found_paragraph:
            for table in doc.tables:
                if para._element.getnext() is table._element:
                    # 获取表格的最右下角单元格内容
                    last_row = table.rows[-1]
                    last_cell = last_row.cells[-1]
                    return last_cell.text.strip()
    return None

# # 使用示例
# csv_path = 'information.csv'  # 替换为你的CSV文件路径


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("f", help="文件路径")
    args = parser.parse_args()
    doc = Document(args.f)

    # 封面日期年月应采用汉字符号书写
    check_date_format(doc)  

    # 设计文件分发表的邮箱、地址、手机号码等应准确，我方人员应为合同制员工
    paragraph_text = '设计文件分发表'  # 替换为段落中的目标文字

    content = find_table_after_paragraph(doc, paragraph_text)
    if content:
        print(content)
    else:
        print(f"未找到位于'{paragraph_text}'段落后面的表格。")
    persons = extract_person_info(content)
    find_person_in_csv("information.csv",persons)
    
    # 目录的字体、字号、行距与正文相同
    check_directory_and_font("xml")   # 等会增加这边的代码

    # 所有正文字体为宋体小四，英文和数字为新罗马，1.5倍行距
    check_paragraph_formatting(doc)

    # 表名置于表上方中央，表号与表的左边线对齐
    check_table_titles(doc) 