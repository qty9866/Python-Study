import win32com.client as win32
from docx import Document
import re
import os


def is_chinese_date_format(date_text):
    pattern = r'[〇一二三四五六七八九十]{4}年[一二三四五六七八九十]{1,2}月'
    return re.match(pattern, date_text) is not None


def mark_incorrect_date(paragraph, incorrect_text):
    start = paragraph.Range.Text.find(incorrect_text)
    rng = paragraph.Range.Duplicate
    rng.Start += start
    rng.End = rng.Start + len(incorrect_text)
    rng.Font.Color = win32.constants.wdColorRed
    paragraph.Range.Comments.Add(rng, "日期格式错误")


def date_format_check(doc):
    for paragraph in doc.Paragraphs:
        text = paragraph.Range.Text.strip()
        match = re.search(r'\d{4}年\d{1,2}月|\d{4}年[一二三四五六七八九十]{1,2}月|[〇一二三四五六七八九十]{4}年\d{1,2}月', text)
        if match:
            date = match.group(0)
            if not is_chinese_date_format(date):
                # 检查是否已经有批注
                if not any(comment.Range.Text == "日期格式错误" for comment in paragraph.Range.Comments):
                    mark_incorrect_date(paragraph, date)
            break


def extract_person_info(docx_path):
    docx = Document(docx_path)
    for para in docx.paragraphs:
        if '设计文件分发表' in para.text:
            for table in docx.tables:
                if para._element.getnext() is table._element:
                    last_row = table.rows[-1]
                    last_cell = last_row.cells[-1]
                    text = last_cell.text.strip()
                    zong_name = re.search(r'项目总负责人：([^\n]+)', text)
                    dan_name = re.search(r'单项设计负责人：([^\n]+)', text)
                    jian_name = re.search(r'建设单位联系人：([^\n]+)', text)
                    tel_match = re.findall(r'电话：([^\n]+)', text)
                    mail_match = re.findall(r'电子邮箱：([^\n]+)', text)
                    persons = []
                    if zong_name and len(tel_match) > 0 and len(mail_match) > 0:
                        persons.append({
                            'name': zong_name.group(1).strip(),
                            'tel': tel_match[0].strip(),
                            'mail': mail_match[0].strip()
                        })
                    if dan_name and len(tel_match) > 1 and len(mail_match) > 1:
                        persons.append({
                            'name': dan_name.group(1).strip(),
                            'tel': tel_match[1].strip(),
                            'mail': mail_match[1].strip()
                        })
                    if jian_name and len(tel_match) > 2 and len(mail_match) > 2:
                        persons.append({
                            'name': jian_name.group(1).strip(),
                            'tel': tel_match[2].strip(),
                            'mail': mail_match[2].strip()
                        })
                    return persons if persons else None
    return None


def check_toc_format(doc):
    if doc.TablesOfContents.Count > 0:
        toc = doc.TablesOfContents(1)
        for para in toc.Range.Paragraphs:
            issues = check_paragraph_format(para, False)
            if issues:
                add_comment(para, issues, "目录")
                break


def check_paragraph_format(paragraph, flag):
    format_issues = set()
    if flag:
        font_name = paragraph.Range.Font.Name
        font_size = paragraph.Range.Font.Size
        if font_name != '宋体' and font_size != 12:
            format_issues.add(f"字体或大小不符合要求")
    if paragraph.LineSpacingRule != win32.constants.wdLineSpace1pt5:
        format_issues.add("行距不为1.5倍")
    return list(format_issues) if format_issues else None


def check_normal_format(doc):
    for para in doc.Paragraphs:
        if not para.Range.Text.strip():
            continue
        if para.Range.Information(12) or para.Range.StoryType in [6, 7]:
            continue
        if para.Style.NameLocal == '正文':
            issues = check_paragraph_format(para, True)
            if issues:
                add_comment(para, issues, "正文")


def check_table_paragraph_alignment(doc):
    for table in doc.Tables:
        if table.Rows.Count == 0:  # 检查是否为空表格
            continue

        # 找到表格前的段落
        prev_range = table.Range.Previous()  # 获取表格前的 Range
        if prev_range is None or prev_range.Paragraphs.Count == 0:
            continue

        prev_paragraph = prev_range.Paragraphs(prev_range.Paragraphs.Count)
        paragraph_text = prev_paragraph.Range.Text.strip()

        # 如果内容为空或为 "设计文件分发表" 则跳过检查
        if not paragraph_text or "设计文件分发表" in paragraph_text:
            continue

        # 检查对齐方式，是否为左对齐
        if prev_paragraph.Alignment != win32.constants.wdAlignParagraphLeft:
            add_comment(prev_paragraph, "此段落未左对齐","表格名")

def add_comment(paragraph, issues, part):
    # 如果段落已有批注，则跳过
    if paragraph.Range.Comments.Count > 0:
        return
    rng = paragraph.Range.Duplicate
    comment_text = part + "未满足的格式要求：" + "；".join(issues)
    paragraph.Range.Comments.Add(rng, comment_text)


if __name__ == '__main__':
    word_app = win32.gencache.EnsureDispatch('Word.Application')
    word_app.Visible = False

    doc_path = 'test.docx'
    doc = word_app.Documents.Open(os.path.abspath(doc_path))

    date_format_check(doc)
    check_toc_format(doc)
    persons = extract_person_info(doc_path)
    print(persons)
    check_table_paragraph_alignment(doc)

    check_normal_format(doc)
    doc.SaveAs(os.path.abspath('test_processed.docx'))
    doc.Close()
    word_app.Quit()
