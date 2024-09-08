import win32com.client as win32
from docx import Document
import re
import os


def is_chinese_date_format(date_text):
    pattern = r'[〇一二三四五六七八九十]{4}年[一二三四五六七八九十]{1,2}月'
    match = re.match(pattern, date_text)
    return match is not None


def mark_incorrect_date(paragraph, incorrect_text):
    start = paragraph.Range.Text.find(incorrect_text)
    # end = start + len(incorrect_text)
    rng = paragraph.Range.Duplicate
    rng.Start += start
    rng.End = rng.Start + len(incorrect_text)
    rng.Font.Color = win32.constants.wdColorRed
    comment = paragraph.Range.Comments.Add(rng, "日期格式错误")
    return comment


# 首页日期格式检查
def date_format_check(doc):
    first_page = doc.Paragraphs
    for paragraph in first_page:
        text = paragraph.Range.Text.strip()
        match = re.search(r'\d{4}年\d{1,2}月|\d{4}年[一二三四五六七八九十]{1,2}月|[〇一二三四五六七八九十]{4}年\d{1,2}月', text)
        if match:
            date = match.group(0)
            if not is_chinese_date_format(date):
                mark_incorrect_date(paragraph, date)
            break


# 获取设计文件分发表中人员信息
def extract_person_info(docx_path):
    docx = Document(docx_path)
    paragraph_text = '设计文件分发表'
    found_paragraph = False
    for para in docx.paragraphs:
        if paragraph_text in para.text:
            for table in docx.tables:
                if para._element.getnext() is table._element:
                    last_row = table.rows[-1]
                    last_cell = last_row.cells[-1]
                    text = last_cell.text.strip()
                    zong_name_match = re.search(r'项目总负责人：([^\n]+)', text)
                    dan_name_match = re.search(r'单项设计负责人：([^\n]+)', text)
                    jian_name_match = re.search(r'建设单位联系人：([^\n]+)', text)
                    tel_match = re.findall(r'电话：([^\n]+)', text)
                    mail_match = re.findall(r'电子邮箱：([^\n]+)', text)
                    persons = []
                    if zong_name_match and len(tel_match) > 0 and len(mail_match) > 0:
                        person1 = {
                            'name': zong_name_match.group(1).strip(),
                            'tel': tel_match[0].strip(),
                            'mail': mail_match[0].strip()
                        }
                        persons.append(person1)
                    if dan_name_match and len(tel_match) > 1 and len(mail_match) > 1:
                        person2 = {
                            'name': dan_name_match.group(1).strip(),
                            'tel': tel_match[1].strip(),
                            'mail': mail_match[1].strip()
                        }
                        persons.append(person2)
                    if jian_name_match and len(tel_match) > 2 and len(mail_match) > 2:
                        person3 = {
                            'name': jian_name_match.group(1).strip(),
                            'tel': tel_match[2].strip(),
                            'mail': mail_match[2].strip()
                        }
                        persons.append(person3)
                    return persons if persons else None
    print("未找到设计文件分发表")
    return None

# 检查文件目录格式字体、字号、行距
def check_toc_format(doc):
    if doc.TablesOfContents.Count > 0:
        toc = doc.TablesOfContents(1)
        for para in toc.Range.Paragraphs:
            issues = check_paragraph_format(para,False)
            if issues:
                add_comment(para, issues,"目录")
                break
    else:
        print("文档中未找到目录")

# 检查正文格式字体为宋体小四，英文和数字为新罗马，1.5倍行距
def check_paragraph_format(paragraph,flag):
    format_issues = set()
    if flag:
        for char in paragraph.Range.Characters:
            font_name = char.Font.Name
            font_size = char.Font.Size
            text = char.Text

            if '\u4e00' <= text <= '\u9fff':  # 中文字符范围
                if font_name != '宋体':
                    format_issues.add("中文字体不是宋体")
                if font_size != 12:  # 小四对应的大小
                    format_issues.add("中文字体大小不是小四")

            # 判断英文字符和数字的字体
            elif text.isalpha() or text.isdigit():  # 英文或数字
                if font_name != 'Times New Roman':
                    format_issues.add("英文或数字字体不是Times New Roman")
                if font_size != 12:
                    format_issues.add("英文或数字字体大小不是小四")
    else:
        font_name = paragraph.Range.Font.Name
        font_size = paragraph.Range.Font.Size

        if font_name != '宋体':
           format_issues.add("字体不是宋体")
        if font_size != 12:
           format_issues.add("字体大小不是小四")
    # 检查段落的行距是否为1.5倍行距
    if paragraph.LineSpacingRule != win32.constants.wdLineSpace1pt5:
        format_issues.add("行距不为1.5倍")

    # 遍历段落中的每个字符，检查字体和字号
    return list(format_issues) if format_issues else None

def check_normal_format(doc):
    # 跳过空的段落
    for para in doc.Paragraphs:
        if not para.Range.Text.strip():
            continue
    # 检查段落是否在表格内，或者是否在页眉、页脚中
        if (para.Range.Information(12) or  # wdWithInTable = 12
                para.Range.StoryType in [6, 7]):  # 使用常量值6和7替代wdHeaderStory和wdFooterStory
            continue  # 跳过表格、页眉和页脚
        if para.Style.NameLocal == '正文':
            issues = check_paragraph_format(para, True)
            if issues:
                add_comment(para, issues,"正文")



def add_comment(paragraph, issues ,part):
    rng = paragraph.Range.Duplicate
    comment_text = part + "未满足的格式要求：" + "；".join(issues)
    paragraph.Range.Comments.Add(rng, comment_text)


if __name__ == '__main__':
    # parser = argparse.ArgumentParser()
    # parser.add_argument('f', help='word file path')
    # args = parser.parse_args()

    word_app = win32.gencache.EnsureDispatch('Word.Application')
    word_app.Visible = False  # 不显示 Word 窗口

    doc_path = 'test.docx'

    doc = word_app.Documents.Open(os.path.abspath(doc_path))

    date_format_check(doc)     # 首页日期格式检查
    check_toc_format(doc)      # 目录格式检查
    persons = extract_person_info(doc_path)     # 提取分发表人员信息
    print(persons)

    check_normal_format(doc)    # 正文格式检测
    doc.SaveAs(os.path.abspath('test_processed.docx'))
    doc.Close()
    word_app.Quit()
